"""
Parser Module
--------------
Extracts raw text and tables from documents.
Routes each file type to the appropriate parsing strategy:
  - PDF → PyMuPDF (fitz) + OCR fallback
  - Images → Tesseract OCR
  - DOCX → python-docx
  - XLSX/CSV → openpyxl / csv
  - HTML → BeautifulSoup
  - TXT/MD → direct read
"""
from __future__ import annotations

import csv
import io
from pathlib import Path

from pipeline.ingestion import DocumentPayload


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def parse(payload: DocumentPayload) -> DocumentPayload:
    """
    Extract raw text and tables from the document in *payload*.
    Mutates and returns the same payload object.
    """
    ext = payload.file_extension

    parser_map = {
        ".pdf": _parse_pdf,
        ".png": _parse_image,
        ".jpg": _parse_image,
        ".jpeg": _parse_image,
        ".tiff": _parse_image,
        ".bmp": _parse_image,
        ".docx": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".csv": _parse_csv,
        ".html": _parse_html,
        ".txt": _parse_text,
        ".md": _parse_text,
    }

    parser_fn = parser_map.get(ext)
    if parser_fn is None:
        raise ValueError(f"[Parser] No parser registered for extension: {ext}")

    parser_fn(payload)

    # Quick sanity log
    text_len = len(payload.raw_text)
    table_count = len(payload.tables)
    print(
        f"[Parser] ✔ Extracted {text_len:,} chars of text, "
        f"{table_count} table(s) from '{payload.file_name}'"
    )
    return payload


# ---------------------------------------------------------------------------
# Individual parsers
# ---------------------------------------------------------------------------

def _parse_pdf(payload: DocumentPayload) -> None:
    """Extract text (and optionally tables) from PDF using PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(payload.file_path)
    pages_text: list[str] = []

    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text)
        else:
            # Fallback: OCR the page image
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            ocr_text = _ocr_bytes(img_bytes)
            pages_text.append(ocr_text)

    doc.close()
    payload.raw_text = "\n\n".join(pages_text)

    # Attempt table extraction using fitz built-in (PyMuPDF ≥ 1.23)
    try:
        doc2 = fitz.open(payload.file_path)
        for page in doc2:
            tabs = page.find_tables()
            for tab in tabs:
                table_data = tab.extract()  # list[list[str]]
                if table_data:
                    payload.tables.append(table_data)
        doc2.close()
    except Exception:
        pass  # Table extraction is best-effort


def _parse_image(payload: DocumentPayload) -> None:
    """OCR an image file."""
    with open(payload.file_path, "rb") as f:
        payload.raw_text = _ocr_bytes(f.read())


def _parse_docx(payload: DocumentPayload) -> None:
    """Extract paragraphs and tables from a Word document."""
    from docx import Document

    doc = Document(payload.file_path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    payload.raw_text = "\n".join(paragraphs)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            payload.tables.append(rows)


def _parse_xlsx(payload: DocumentPayload) -> None:
    """Extract text and table data from an Excel workbook."""
    from openpyxl import load_workbook

    wb = load_workbook(payload.file_path, data_only=True)
    all_text: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        table_rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            str_row = [str(c) if c is not None else "" for c in row]
            table_rows.append(str_row)
            all_text.append(" | ".join(str_row))
        if table_rows:
            payload.tables.append(table_rows)

    payload.raw_text = "\n".join(all_text)


def _parse_csv(payload: DocumentPayload) -> None:
    """Parse a CSV file into text and a table."""
    with open(payload.file_path, newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        rows = [row for row in reader]

    if rows:
        payload.tables.append(rows)

    payload.raw_text = "\n".join([", ".join(r) for r in rows])


def _parse_html(payload: DocumentPayload) -> None:
    """Extract visible text from HTML."""
    from bs4 import BeautifulSoup

    with open(payload.file_path, encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")

    # Remove scripts and styles
    for tag in soup(["script", "style"]):
        tag.decompose()

    payload.raw_text = soup.get_text(separator="\n", strip=True)

    # Extract HTML tables
    for html_table in soup.find_all("table"):
        rows = []
        for tr in html_table.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            rows.append(cells)
        if rows:
            payload.tables.append(rows)


def _parse_text(payload: DocumentPayload) -> None:
    """Read plain text or markdown files."""
    with open(payload.file_path, encoding="utf-8") as f:
        payload.raw_text = f.read()


# ---------------------------------------------------------------------------
# OCR helper
# ---------------------------------------------------------------------------

def _ocr_bytes(image_bytes: bytes) -> str:
    """Run Tesseract OCR on raw image bytes."""
    from PIL import Image
    import pytesseract

    img = Image.open(io.BytesIO(image_bytes))
    text: str = pytesseract.image_to_string(img)
    return text
